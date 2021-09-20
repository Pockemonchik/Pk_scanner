import winapps
import psutil
import os
import platform
import sys
import GPUtil
import wmi
from psutil import virtual_memory
from browser_history.browsers import Firefox
from winreg import *
from docx import Document
from PyQt5 import QtWidgets,QtGui
from design import Ui_MainWindow
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


w=wmi.WMI()

def app_list():
    """Поиск спсика установленных программ с версиями"""
    app_list = []
    for app in winapps.list_installed():
        app_list.append((app.name, app.version))
    return app_list

def pk_info():
   
    mem = virtual_memory()
    result = dict()
    """Получаем информацию об ОС, процессоре, РАМ"""
    result["os"] = platform.uname().system, platform.uname().version
    result["cpu"] = platform.processor()
    result["ram"] = (mem.total)//1000000000

    """Получаем видеопамять"""
    gpus = GPUtil.getGPUs()
    gpus_list = []
    gpu_total_memory = 0
    for gpu in gpus:
        
        gpu_total_memory += gpu.memoryTotal

       
    print ("Видеопамять")
    result["video"] = gpu_total_memory

    """Обьем диска"""
    obj_Disk = psutil.disk_usage('/')

    result["pzu"] = obj_Disk.total // (1024.0 ** 3)
    return result

def user_info():
    """Получаем список учеток"""

    for u in w.Win32_UserAccount(["Name"]): #Net
        print (u.Name)

def search_USB():
    REG_PATH = r"SYSTEM\\CurrentControlSet\\Enum\\USBSTOR\\"
    try:
        registry_key = OpenKey(HKEY_LOCAL_MACHINE, REG_PATH, 0,
                                       KEY_READ) 
        usb_list = []
        id_list = []
        i = 0
        while True:
            try:
                # смотрим ключи в usbstore
                name = EnumKey(registry_key, i)
                usb_key = OpenKey(HKEY_LOCAL_MACHINE, r"SYSTEM\\CurrentControlSet\\Enum\\USBSTOR\\"+name, 0,
                                       KEY_READ) 
                # получаем id в подкаталоге
                j = 0
                while True:
                    try:
                        
                        usb_name = OpenKey(HKEY_LOCAL_MACHINE, r"SYSTEM\\CurrentControlSet\\Enum\\USBSTOR\\"+name+"\\"+EnumKey(usb_key,j), 0,
                                       KEY_READ)
                                       
                        usb_list.append((QueryValueEx(usb_name,"FriendlyName")[0],EnumKey(usb_key,j)))
                        j +=1
                    except OSError as e:
                        if isinstance(e, WindowsError) and e.winerror == 259:
                            break
                        else:
                            raise e

                usb_key.Close()
                i+= 1
            except OSError as e:
                if isinstance(e, WindowsError) and e.winerror == 259:
                    break
                else:
                    raise e
        # value, regtype = QueryValueEx(registry_key, name)
        CloseKey(registry_key)
        return usb_list
        # return value
    except WindowsError:
        return None

def browser_history_list():
    """Получаем истою браузера"""
    f = Firefox()
    outputs = f.fetch_history()
    his = outputs.histories
    result = []
    for url in his:
        result.append((url[0].date(),url[1]))
    print("история браузера")
    return result

def network_profiles():
    REG_PATH = r"SOFTWARE\\Microsoft\\Windows NT\\CurrentVersion\\NetworkList\\Profiles\\"
    network_profiles = []
    try:
        registry_key = OpenKey(HKEY_LOCAL_MACHINE, REG_PATH, 0,
                                       KEY_READ) 
        i = 0
        while True:
            try:
                name = EnumKey(registry_key, i)
                i+= 1
                profile_key = OpenKey(HKEY_LOCAL_MACHINE, r"SOFTWARE\\Microsoft\\Windows NT\\CurrentVersion\\NetworkList\\Profiles\\"+name, 0,
                                       KEY_READ) 
                # получаем id в подкаталоге
                try:
                    network_profiles.append(QueryValueEx(profile_key,"ProfileName"))
                except OSError as e:
                    if isinstance(e, WindowsError) and e.winerror == 259:
                        network_profiles.append("Не удалось получить данные, запустите программу от имени администратора")
                        break
                    else:
                        raise e

                profile_key.Close()
            except OSError as e:
                if isinstance(e, WindowsError) and e.winerror == 259:
                    network_profiles.append("Не удалось получить данные, запустите программу от имени администратора")
                    break
                else:
                    raise e
        CloseKey(registry_key)
        return network_profiles
        # return value
    except WindowsError as e:
        network_profiles.append("Не удалось получить данные, запустите программу от имени администратора")
        print(e)
        return network_profiles

def create_report(fio,service,app_list,pk_info,network_profiles,browser_history_list,search_USB,path_to_document):
    """Заполняем отчет"""
    document = Document("report.docx")
    #сотрудник
    document.paragraphs[1].text+=fio
    document.paragraphs[2].text+=service

    #конфигурация АРм
    document.tables[0].rows[1].cells[0].text = pk_info["os"]
    document.tables[0].rows[1].cells[1].text = pk_info["cpu"]
    document.tables[0].rows[1].cells[2].text = str(pk_info["ram"])
    document.tables[0].rows[1].cells[3].text = str(pk_info["video"])
    document.tables[0].rows[1].cells[4].text = str(pk_info["pzu"])   
    if pk_info["ram"]<2:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#cd0221"/>'.format(nsdecls('w')))
        document.tables[0].rows[1].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
    if pk_info["video"]<512:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#cd0221"/>'.format(nsdecls('w')))
        document.tables[0].rows[1].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
    if pk_info["pzu"]<100:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#cd0221"/>'.format(nsdecls('w')))
        document.tables[0].rows[1].cells[4]._tc.get_or_add_tcPr().append(shading_elm_1)
    
    # ПО и СЗИ 
    szi_list = ["Kaspersky","sudis","secret","csp","vip","office"]
    szi_count = 1
    for app in app_list:
        row = document.tables[1].add_row()
        row.cells[0].text = app[0]
        row.cells[1].text = str(app[1])
        for szi in szi_list:
            if szi.lower() in app[0].lower():
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#419248"/>'.format(nsdecls('w')))
                document.tables[1].rows[szi_count].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        szi_count+=1
    
    # Сетевая активность
    print(network_profiles)
    for profile in network_profiles:
        row = document.tables[2].add_row()
        row.cells[0].text = str(profile[0])

    for href in browser_history_list:
        row = document.tables[3].add_row()
        row.cells[0].text = str(href[1])
        row.cells[1].text = str(href[0])

    # Носители информации
    for usb in search_USB:
        row = document.tables[4].add_row()
        row.cells[0].text = usb[0]
        row.cells[1].text = usb[1]
        


    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    for paragraph in document.paragraphs:
        paragraph.style = document.styles['Normal']
    
    document.save(path_to_document+" "+service+"_"+fio+".docx")


class mywindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(mywindow, self).__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda:self._get_report("_"))
        self.ui.pushButton_2.clicked.connect(self.saveFileDialog)
        self.ui.progressBar.setValue(0)
        
    

    def _get_report(self,path_to_document):
        self.ui.label_3.setText("Загрузка")
        apps = app_list()
        pk = pk_info()
        self.ui.progressBar.setValue(0)
        net = network_profiles()
        self.ui.progressBar.setValue(25)
        browser = browser_history_list()
        self.ui.progressBar.setValue(50)
        usb = search_USB()
        self.ui.progressBar.setValue(75)
        create_report(self.ui.lineEdit.text(),self.ui.lineEdit_2.text(),apps,pk,net,browser,usb,path_to_document)
        self.ui.progressBar.setValue(100)
        self.ui.label_3.setText("Файл успешно сохранен")
    
    def saveFileDialog(self):
        options = QtWidgets.QFileDialog.Options()
        options |= QtWidgets.QFileDialog.DontUseNativeDialog
        fileName, _ = QtWidgets.QFileDialog.getSaveFileName(self,"Сохраненния отчета","","All Files (*);;Text Files (*.txt)", options=options)
        if fileName != "" and fileName != None:
            print(fileName,"filename dialog")
            self._get_report(fileName)
            


app = QtWidgets.QApplication([])
application = mywindow()

application.show()
 
sys.exit(app.exec())
