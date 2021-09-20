import re
import psutil
import os
import platform
import gzip 
import io
import sys
import GPUtil
import glob
from psutil import virtual_memory
from docx import Document
from PyQt5 import QtWidgets,QtGui
from design import Ui_MainWindow
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import datetime as DT
import os
import sqlite3
from pathlib import Path




def app_list():
    """Поиск спсика установленных программ с версиями"""
    app_list = []
    filter = ["libreoff","kasper","sudis","vip"]
    dpkg_out = []
    for f in filter:
        dpkg_out.append(os.popen("dpkg -l |grep "+f.lower()).read().split())
    for dpkg in dpkg_out:
        for i,pac in enumerate(dpkg):
            if pac == "ii":
                app_list.append((dpkg[i+1],dpkg[i+2]))
    return app_list

def pk_info():
    mem = virtual_memory()
    result = dict()
    """Получаем информацию об ОС, процессоре, РАМ"""
    result["os"] = platform.uname().system, platform.uname().version
    # result["cpu"] = subprocess.call(["cat ","/proc/cpuinfo "," |", "grep", "name","|", "uniq"])
    result["cpu"] = os.popen("cat /proc/cpuinfo  | grep 'name'| uniq").read()

    result["ram"] = (mem.total)//1000000000

    """Получаем видеопамять"""
    gpus = GPUtil.getGPUs()
    gpus_list = []
    gpu_total_memory = 0
    for gpu in gpus:
        
        gpu_total_memory += gpu.memoryTotal
        
    result["video"] = gpu_total_memory

    """Обьем диска"""
    obj_Disk = psutil.disk_usage('/')

    result["pzu"] = obj_Disk.total // (1024.0 ** 3)
    return result

def user_info():
    """Получаем список учеток"""
    user_info = []
    return user_info

def search_USB():
    usblist = []
    data = []
    log_dir = os.listdir(path='/var/log/')
    for log_name in log_dir:
        if "syslog" in log_name and "gz" in log_name:
            with gzip.open("/var/log/"+log_name, 'rb') as input_file: 
                with io.TextIOWrapper(input_file, encoding='utf-8') as dec:
                    zip_data = dec.read().split("\n")
                    for i,item in enumerate(zip_data):
                        if "usb" in item and "SerialNumber:" in item:
                            data.append((item,zip_data[i-2],zip_data[i-1])) 
        if "syslog" in log_name and "gz" not in log_name:
            with open("/var/log/"+log_name, 'rb') as input_file: 
                with io.TextIOWrapper(input_file, encoding='utf-8') as dec:
                    zip_data = dec.read().split("\n")
                    for i,item in enumerate(zip_data):
                        if "usb" in item and "SerialNumber:" in item:
                            data.append((item,zip_data[i-2],zip_data[i-1])) 
    print(data)
    for usb in data:
        usblist.append((re.findall(r"(Manufacturer:[\s\S]+)",usb[2])[0].split(" ",1)[1]
        +" "+re.findall(r"(Product:[\s\S]+)",usb[1])[0].split(" ",1)[1],
        re.findall(r"(SerialNumber:[\s\S]+)",usb[0])[0].split(" ",1)[1]))

    return usblist

def browser_history_list():
    """Получаем истою браузера"""    
    result =[]
    try:
        history_files = glob.glob("/home/andrey/.mozilla/firefox/*.default*",recursive=False)
        print(history_files)
        for file in history_files:
            try:
                db = sqlite3.connect(file+"/places.sqlite")
                urls = db.execute('select url from moz_places')
                for [url] in urls:
                    print(url)
                    result.append(url)
            except Exception as e:
                print(e)
                continue
    except (IndexError, sqlite3.OperationalError) as e:
        print(e)
        sys.exit(__doc__)  # print usage and exit

    print("история браузера")
    return result

def printer_list():
    printer_list = []
    data = os.popen("lpstat -W completed -u $(getent passwd | awk -F: '{print $1}' | paste -sd ',')").read()
    printer_list = data.split("\n")
    return printer_list
       

def create_report(fio,service,app_list,pk_info,printer_list,browser_history_list,search_USB,path_to_document):
    try:
        """Заполняем отчет"""
        document = Document("report.docx")
        #сотрудник
        document.paragraphs[1].text+=fio
        document.paragraphs[2].text+=service

        #конфигурация АРм
        document.tables[0].rows[1].cells[0].text = pk_info["os"]
        document.tables[0].rows[1].cells[1].text = pk_info["cpu"]
        document.tables[0].rows[1].cells[2].text = str(pk_info["ram"])
        document.tables[0].rows[1].cells[3].text = str(pk_info["video"])
        document.tables[0].rows[1].cells[4].text = str(pk_info["pzu"])   
        if pk_info["ram"]<2:
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#cd0221"/>'.format(nsdecls('w')))
            document.tables[0].rows[1].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
        if pk_info["video"]<512:
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#cd0221"/>'.format(nsdecls('w')))
            document.tables[0].rows[1].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)
        if pk_info["pzu"]<100:
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#cd0221"/>'.format(nsdecls('w')))
            document.tables[0].rows[1].cells[4]._tc.get_or_add_tcPr().append(shading_elm_1)
        
        # ПО и СЗИ 
        szi_list = ["Kaspersky","sudis","secret","csp","vip","android","office"]
        szi_count = 1
        for app in app_list:
            row = document.tables[1].add_row()
            row.cells[0].text = app[0]
            row.cells[1].text = str(app[1])
            for szi in szi_list:
                if szi.lower() in app[0].lower():
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#419248"/>'.format(nsdecls('w')))
                    document.tables[1].rows[szi_count].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
            szi_count+=1
        
        # Сетевая активность and printers
        for printer in printer_list:
            row = document.tables[2].add_row()
            row.cells[0].text = str(printer)

        for i,href in enumerate(browser_history_list):
            row = document.tables[3].add_row()
            row.cells[0].text = str(href)

            
            if "mvd" in str(href).lower():
                print(href,"href")
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#419248"/>'.format(nsdecls('w')))
                document.tables[3].rows[i+1].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

        

        # Носители информации
        print("USB")
        print(search_USB)
        for usb in search_USB:
            row = document.tables[4].add_row()
            row.cells[0].text = usb[0]
            row.cells[1].text = usb[1]
            


        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']
        
        document.save(path_to_document+" "+service+"_"+fio+".docx")
        return 0
    except Exception as e:
        print(e)
        return e
        

class mywindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(mywindow, self).__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(lambda:self._get_report("_"))
        self.ui.pushButton_2.clicked.connect(self.saveFileDialog)
        self.ui.progressBar.setValue(0)
        
    

    def _get_report(self,path_to_document):
        self.ui.label_3.setText("Загрузка")
        apps = app_list()
        pk = pk_info()
        self.ui.progressBar.setValue(0)
        printers = printer_list()
        self.ui.progressBar.setValue(25)
        browser = browser_history_list()
        self.ui.progressBar.setValue(50)
        usb = search_USB()
        self.ui.progressBar.setValue(75)
        result = create_report(self.ui.lineEdit.text(),self.ui.lineEdit_2.text(),apps,pk,printers,browser,usb,path_to_document)
        self.ui.progressBar.setValue(100)
        if result == 0:
            self.ui.label_3.setText("Файл успешно сохранен")
        else:
            self.ui.label_3.setText("Error!"+ str(result))
        
    
    def saveFileDialog(self):
        options = QtWidgets.QFileDialog.Options()
        options |= QtWidgets.QFileDialog.DontUseNativeDialog
        fileName, _ = QtWidgets.QFileDialog.getSaveFileName(self,"Сохраненния отчета","","All Files (*);;Text Files (*.txt)", options=options)
        if fileName != "" and fileName != None:
            print(fileName,"filename dialog")
            self._get_report(fileName)


            


app = QtWidgets.QApplication([])
application = mywindow()

application.show()
 
sys.exit(app.exec())
